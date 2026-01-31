import streamlit as st
import os
from groq import Groq
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from fpdf import FPDF
import io
import re
import json
from pypdf import PdfReader
import requests # عشان نحمل الخط
import arabic_reshaper # عشان يشبك الحروف
from bidi.algorithm import get_display # عشان يظبط الاتجاه من اليمين للشمال

# --- 1. Page Configuration ---
st.set_page_config(
    page_title="Elite CV Builder",
    page_icon="🚀",
    layout="wide",
    initial_sidebar_state="expanded"
)

# --- 2. Auto-Download Arabic Font ---
# بنحمل خط "Amiri" عشان هو ممتاز في العربي والإنجليزي مع بعض
FONT_URL = "https://github.com/google/fonts/raw/main/ofl/amiri/Amiri-Regular.ttf"
FONT_PATH = "Amiri-Regular.ttf"

def check_and_download_font():
    if not os.path.exists(FONT_PATH):
        with st.spinner("Downloading Arabic Font support..."):
            response = requests.get(FONT_URL)
            with open(FONT_PATH, "wb") as f:
                f.write(response.content)

check_and_download_font()

# --- 3. API & Sidebar Configuration ---
api_key = None
using_shared_key = False

with st.sidebar:
    st.image("https://cdn-icons-png.flaticon.com/512/3135/3135715.png", width=50)
    st.title("💡 Quick Guide")
    
    st.markdown("""
    **How to build a world-class CV:**
    1. **Have an old CV?** Upload it in Step 1 to auto-fill.
    2. **Writer's Block?** Use "Get Suggestions" in Step 3.
    3. **Finish:** Download your ATS-optimized Resume.
    """)
    
    st.divider()
    
    with st.expander("⚙️ Advanced Settings"):
        use_own_key = st.checkbox("Use my own API Key", value=False)
        if use_own_key:
            user_input_key = st.text_input("Groq API Key", type="password")
            if user_input_key:
                api_key = user_input_key
                using_shared_key = False
        else:
            if "GROQ_API_KEY" in st.secrets:
                api_key = st.secrets["GROQ_API_KEY"]
                using_shared_key = True
                st.success("✅ Connected to Shared Server")
            else:
                st.warning("⚠️ No API Key found.")

# Validate Connection
if not api_key:
    st.warning("⚠️ Please configure the API Key in the sidebar to proceed.")
    st.stop()

client = Groq(api_key=api_key)
MODEL_NAME = "llama-3.3-70b-versatile"

# --- 4. Helper Functions ---

# دالة لمعالجة النص العربي (تشبيك الحروف + عكس الاتجاه)
def process_text_for_pdf(text):
    if not text: return ""
    try:
        # Reshape: بيخلي الحروف تشبك في بعض (ل -> لـ)
        reshaped_text = arabic_reshaper.reshape(text)
        # Bidi: بيخلي الكلام من اليمين للشمال
        bidi_text = get_display(reshaped_text)
        return bidi_text
    except:
        return text

def extract_text_from_pdf(file):
    reader = PdfReader(file)
    text = ""
    for page in reader.pages:
        text += page.extract_text()
    return text

def extract_text_from_docx(file):
    doc = Document(file)
    return "\n".join([para.text for para in doc.paragraphs])

def parse_resume_with_ai(text):
    prompt = f"""
    Extract the following details from this resume text:
    Name, Email, Phone, City, LinkedIn, Target Job Title (infer if not present), 
    Skills (comma-separated), Experience (raw text),
    University, College (Faculty), Degree, and Graduation Year.
    
    Resume Text:
    {text[:4000]} 
    
    Output ONLY a valid JSON object with these keys: 
    "name", "email", "phone", "city", "linkedin", "target_title", "skills", "experience", 
    "university", "college", "degree", "grad_year".
    """
    try:
        completion = client.chat.completions.create(
            model=MODEL_NAME,
            messages=[{"role": "user", "content": prompt}],
            temperature=0.1,
            response_format={"type": "json_object"}
        )
        return json.loads(completion.choices[0].message.content)
    except Exception as e:
        return None

def get_job_suggestions(role_title):
    prompt = f"""
    Give me 5 professional, metric-driven bullet points for a "{role_title}" resume.
    Write them in English. Start with strong action verbs.
    Output ONLY the bullet points.
    """
    try:
        completion = client.chat.completions.create(
            model=MODEL_NAME,
            messages=[{"role": "user", "content": prompt}],
            temperature=0.7
        )
        return completion.choices[0].message.content
    except:
        return "Error generating suggestions."

def safe_generate(prompt_text):
    try:
        completion = client.chat.completions.create(
            model=MODEL_NAME,
            messages=[
                {"role": "system", "content": "You are a Senior HR Recruiter. Output strict, clean text. Do NOT use markdown bold (**). Do NOT number the sections."},
                {"role": "user", "content": prompt_text}
            ],
            temperature=0.3,
            max_tokens=3500,
        )
        return completion.choices[0].message.content
    except Exception as e:
        return f"Error: {str(e)}"

# --- 5. File Generation Functions ---

def create_docx(text):
    doc = Document()
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5); section.bottom_margin = Inches(0.5); section.left_margin = Inches(0.5); section.right_margin = Inches(0.5)
    
    text = text.replace("**", "").replace("##", "")
    for line in text.split('\n'):
        line = line.strip()
        if not line: continue
        
        line_no_num = re.sub(r'^\d+\.\s*', '', line)
        
        if line_no_num.isupper() and len(line_no_num) < 60 and "|" not in line:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12); p.paragraph_format.space_after = Pt(6)
            run = p.add_run(line_no_num); run.bold = True; run.font.size = Pt(12)
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER if "NAME" not in line else WD_PARAGRAPH_ALIGNMENT.LEFT
        elif "|" in line and "@" in line:
            p = doc.add_paragraph(line); p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER; p.paragraph_format.space_after = Pt(12)
        elif "|" in line and "@" not in line:
            p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(8); run = p.add_run(line); run.bold = True; run.font.size = Pt(11); p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        elif line.startswith('-') or line.startswith('•'):
            clean_line = line.replace('•', '').replace('-', '').strip(); p = doc.add_paragraph(clean_line, style='List Bullet'); p.paragraph_format.space_after = Pt(2) 
        else:
            p = doc.add_paragraph(line); p.paragraph_format.space_after = Pt(2)
    buffer = io.BytesIO(); doc.save(buffer); buffer.seek(0)
    return buffer

# --- UPDATED PDF FUNCTION FOR ARABIC ---
def create_pdf(text):
    class PDF(FPDF):
        def header(self): pass
        def footer(self): pass
    
    pdf = PDF()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=15)
    
    # 1. تسجيل الخط العربي الذي قمنا بتحميله
    # بنسميه 'Amiri' وبنقوله إن ده خط Unicode
    pdf.add_font('Amiri', '', FONT_PATH, uni=True)
    pdf.set_font('Amiri', '', 11) # استخدام الخط كـ Default
    
    text = text.replace("**", "").replace("##", "")
    
    for line in text.split('\n'):
        line = line.strip()
        if not line: continue
        if "___" in line: continue
        
        # معالجة النص العربي قبل الكتابة
        line = process_text_for_pdf(line)
        
        # محاولة اكتشاف العناوين (Headers)
        # بما إننا عملنا Reshape، الـ Upper مش هيشتغل على العربي، بس شغال للإنجليزي
        is_header = False
        # شرط تقريبي للعناوين: لو السطر قصير ومفهوش علامات ترقيم كتير
        if len(line) < 50 and "|" not in line and "." not in line and not line.startswith("-"): 
             # لو إنجليزي وكله كابيتال يعتبر عنوان
             if re.search(r'[A-Z]', line) and line.isupper(): is_header = True
             # لو عربي، صعب نحدد، بس ممكن نعتمد على المكان (حالياً هنخليه Bold بس)

        if is_header:
            pdf.ln(6)
            pdf.set_font("Amiri", '', 13) # خط أكبر للعناوين
            # في الـ PDF العناوين العربي بتطلع محتاجة محاذاة يمين (R) أو وسط (C)
            # بما إننا عملنا Bidi، الكلام العربي هيظهر صح بس المحاذاة مهمة
            pdf.cell(0, 6, line, ln=True, align='C') 
            
            # رسم خط تحت العنوان
            x = pdf.get_x(); y = pdf.get_y()
            pdf.line(10, y, 200, y) 
            pdf.ln(4)
            pdf.set_font("Amiri", '', 11) # نرجع للخط العادي
            
        elif "|" in line and "@" in line: # سطر الكونتاكت
            pdf.set_font("Amiri", '', 10)
            pdf.multi_cell(0, 5, line, align='C')
            pdf.ln(4)
            
        elif "|" in line and "@" not in line: # تفاصيل الشغل (Role | Company)
            pdf.ln(4)
            pdf.set_font("Amiri", '', 11) # ممكن نعمله Bold لو لقينا font bold
            pdf.cell(0, 6, line, ln=True, align='L' if re.search(r'[a-zA-Z]', line) else 'R') # محاذاة حسب اللغة
            pdf.ln(2)
            
        elif line.startswith('-') or line.startswith('•'): # القوائم (Bullets)
            pdf.set_font("Amiri", '', 11)
            # تنظيف الرموز عشان العربي
            clean_line = line.replace('-', '').replace('•', '').strip()
            # إضافة نقطة يدوية عشان Bidi ممكن يعكس مكانها
            pdf.multi_cell(0, 5, "• " + clean_line, align='L' if re.search(r'[a-zA-Z]', clean_line) else 'R')
            pdf.ln(2)
            
        else: # نص عادي
            pdf.set_font("Amiri", '', 11)
            pdf.multi_cell(0, 5, line, align='L' if re.search(r'[a-zA-Z]', line) else 'R')
            pdf.ln(1)
    
    buffer = io.BytesIO()
    # هنا شلنا encode('latin-1') عشان ده اللي كان بيبوظ العربي
    # FPDF with uni=True بيطلع binary مظبوط
    pdf_output = pdf.output(dest='S')
    
    # تحويل الـ string لـ bytes لو لزم الأمر (حسب نسخة fpdf)
    if isinstance(pdf_output, str):
        buffer.write(pdf_output.encode('latin-1')) # fallback for old fpdf if uni=True fails, but usually with font it works differently.
        # CORRECTION: With Unicode font in FPDF, we usually get bytes or strings that handle unicode properly.
        # Let's rely on standard bytes output for streamlit.
        # Actually, let's strictly write bytes:
        pass 
    
    # الطريقة الأضمن مع الخطوط الخارجية:
    buffer.write(pdf.output(dest='S').encode("latin1")) # FPDF 1.7.2 workaround for unicode injection
    # *توضيح*: مع FPDF العادية واستخدام ttf، الدالة output بترجع string بترميز خاص، الـ encode("latin1") هنا مش بيبوظ العربي، هو بيحول الـ binary string لـ bytes عشان الـ Buffer.
    
    buffer.seek(0)
    return buffer

# --- 6. Session State ---
if 'step' not in st.session_state: st.session_state.step = 1
if 'cv_data' not in st.session_state: st.session_state.cv_data = {}
for key in ['final_cv', 'cover_letter', 'ats_analysis']:
    if key not in st.session_state: st.session_state[key] = ""

if st.session_state.step > 6: st.session_state.step = 1; st.rerun()

def next_step(): st.session_state.step += 1
def prev_step(): st.session_state.step -= 1

# --- 7. Main App UI ---
st.title("🚀 Elite CV Builder")
st.markdown("##### Your AI-Powered Assistant for ATS-Optimized Resumes")

if st.session_state.step < 6: st.progress(st.session_state.step / 6)

# STEP 1: Personal Info
if st.session_state.step == 1:
    st.header("1️⃣ Personal Information")
    
    with st.expander("📄 Have an old CV? Upload to Auto-Fill", expanded=False):
        uploaded_file = st.file_uploader("Upload PDF or Word file", type=['pdf', 'docx', 'doc'])
        if uploaded_file is not None:
            if st.button("🧠 Auto-Fill with AI"):
                with st.spinner("Reading file..."):
                    try:
                        if uploaded_file.name.endswith('.pdf'): text = extract_text_from_pdf(uploaded_file)
                        else: text = extract_text_from_docx(uploaded_file)
                        parsed_data = parse_resume_with_ai(text)
                        if parsed_data:
                            st.session_state.cv_data.update(parsed_data)
                            st.success("Data extracted! Please review below.")
                            st.rerun()
                        else: st.error("Could not parse file.")
                    except Exception as e: st.error(f"Error: {e}")

    st.info("Or fill in your details manually (Arabic or English supported):")
    with st.form("step1"):
        col1, col2 = st.columns(2)
        with col1:
            name = st.text_input("Full Name", st.session_state.cv_data.get('name', ''))
            email = st.text_input("Email", st.session_state.cv_data.get('email', ''))
            city = st.text_input("City/Country", st.session_state.cv_data.get('city', ''))
            portfolio = st.text_input("Portfolio Link", st.session_state.cv_data.get('portfolio', ''))
        with col2:
            phone = st.text_input("Phone Number", st.session_state.cv_data.get('phone', ''))
            linkedin = st.text_input("LinkedIn Profile", st.session_state.cv_data.get('linkedin', ''))
            github = st.text_input("GitHub Profile", st.session_state.cv_data.get('github', ''))
        
        st.markdown("---")
        target_title = st.text_input("🔴 Target Job Title (Important for ATS)", st.session_state.cv_data.get('target_title', ''))
        
        st.markdown("### 🎓 Education")
        c1, c2, c3, c4 = st.columns(4)
        with c1: university = st.text_input("University", st.session_state.cv_data.get('university', ''))
        with c2: college = st.text_input("College/Faculty", st.session_state.cv_data.get('college', ''), placeholder="e.g. Faculty of Engineering")
        with c3: degree = st.text_input("Degree", st.session_state.cv_data.get('degree', ''))
        with c4: grad_year = st.text_input("Grad Year", st.session_state.cv_data.get('grad_year', ''))

        if st.form_submit_button("Next Step ➡️"):
            if name and target_title:
                st.session_state.cv_data.update({
                    'name':name, 'email':email, 'phone':phone, 'linkedin':linkedin, 'city':city, 
                    'portfolio':portfolio, 'github':github, 
                    'target_title':target_title, 
                    'university':university, 
                    'college':college, 
                    'degree':degree, 
                    'grad_year':grad_year
                })
                next_step(); st.rerun()
            else: st.warning("Name and Target Job Title are required!")

# STEP 2: Skills
elif st.session_state.step == 2:
    st.header("2️⃣ Skills")
    with st.form("step2"):
        st.write("List your technical and soft skills:")
        skills = st.text_area("Skills", st.session_state.cv_data.get('skills', ''), height=150)
        languages = st.text_input("Languages", st.session_state.cv_data.get('languages', ''))
        
        col1, col2 = st.columns([1, 5])
        with col1: 
            if st.form_submit_button("Back"): prev_step(); st.rerun()
        with col2:
            if st.form_submit_button("Next Step ➡️"):
                st.session_state.cv_data.update({'skills':skills, 'languages':languages})
                next_step(); st.rerun()

# STEP 3: Experience
elif st.session_state.step == 3:
    st.header("3️⃣ Professional Experience")
    
    st.markdown("##### ✨ Get AI Suggestions:")
    c_input, c_btn = st.columns([3, 1])
    with c_input:
        default_role = st.session_state.cv_data.get('target_title', '')
        suggestion_role = st.text_input("Enter Role Title for Suggestions", value=default_role, label_visibility='collapsed')
    with c_btn:
        if st.button("Get Suggestions 🧠", use_container_width=True):
            if suggestion_role:
                with st.spinner("Thinking..."):
                    sugg = get_job_suggestions(suggestion_role)
                    current_text = st.session_state.cv_data.get('raw_experience', '')
                    st.session_state.cv_data['raw_experience'] = current_text + "\n" + sugg
                    st.rerun()
            else: st.warning("Please enter a role title!")

    with st.form("step3"):
        st.write("👇 Edit your experience here:")
        raw_experience = st.text_area("Experience:", st.session_state.cv_data.get('raw_experience', ''), height=250)
        
        col1, col2 = st.columns([1, 5])
        with col1: 
            if st.form_submit_button("Back"): prev_step(); st.rerun()
        with col2:
            if st.form_submit_button("Next Step ➡️"):
                st.session_state.cv_data['raw_experience'] = raw_experience
                next_step(); st.rerun()

# STEP 4: Projects
elif st.session_state.step == 4:
    st.header("4️⃣ Projects & Certifications")
    with st.form("step4"):
        projects = st.text_area("Projects:", st.session_state.cv_data.get('projects', ''))
        certs = st.text_area("Certifications:", st.session_state.cv_data.get('certs', ''))
        volunteering = st.text_area("Volunteering:", st.session_state.cv_data.get('volunteering', ''))
        
        col1, col2 = st.columns([1, 5])
        with col1: 
            if st.form_submit_button("Back"): prev_step(); st.rerun()
        with col2:
            if st.form_submit_button("Next Step ➡️"):
                st.session_state.cv_data.update({'projects':projects, 'certs':certs, 'volunteering':volunteering})
                next_step(); st.rerun()

# STEP 5: Target Job
elif st.session_state.step == 5:
    st.header("5️⃣ Target Job Details")
    with st.form("step5"):
        st.write("Paste the Job Description (JD):")
        target_job = st.text_area("Job Description:", st.session_state.cv_data.get('target_job', ''), height=150)
        
        col1, col2 = st.columns([1, 5])
        with col1: 
            if st.form_submit_button("Back"): prev_step(); st.rerun()
        with col2:
            if st.form_submit_button("🚀 Generate CV"):
                st.session_state.cv_data['target_job'] = target_job
                next_step(); st.rerun()

# STEP 6: Result
elif st.session_state.step == 6:
    st.balloons()
    st.success("🎉 Congratulations! Your CV is ready.")
    
    safe_name = re.sub(r'[^a-zA-Z0-9]', '_', st.session_state.cv_data.get('name', 'User'))
    file_name = f"{safe_name}_CV.pdf"
    word_file_name = f"{safe_name}_CV.docx"

    t1, t2, t3 = st.tabs(["📄 Resume Preview", "✉️ Cover Letter", "📊 ATS Score"])
    jd = st.session_state.cv_data.get('target_job', '')

    with t1:
        if not st.session_state.final_cv:
            with st.spinner("⏳ Writing your resume..."):
                contact_parts = [st.session_state.cv_data[k] for k in ['phone', 'city', 'email', 'linkedin', 'github', 'portfolio'] if st.session_state.cv_data.get(k)]
                contact_line = " | ".join(contact_parts)

                optional_prompt = ""
                if st.session_state.cv_data.get('projects'): optional_prompt += f"\n5. **PROJECTS**\n   - {st.session_state.cv_data['projects']}"
                if st.session_state.cv_data.get('certs'): optional_prompt += f"\n6. **CERTIFICATIONS**\n   - {st.session_state.cv_data['certs']}"
                if st.session_state.cv_data.get('volunteering'): optional_prompt += f"\n7. **VOLUNTEERING**\n   - {st.session_state.cv_data['volunteering']}"

                prompt_cv = f"""
                Act as a Senior Resume Expert. Write a professional CV based on this data.
                **RULES:**
                1. Clean Text Only (No markdown bold).
                2. No Section Numbers.
                3. Keep the input language (If user wrote in Arabic, keep it in Arabic). 
                4. Dates: Use "Mon YYYY" format.
                
                **HEADER:**
                {st.session_state.cv_data['name'].upper()}
                {contact_line}
                
                **SECTIONS:**
                PROFESSIONAL SUMMARY (Tailored to {st.session_state.cv_data['target_title']})
                TECHNICAL SKILLS ({st.session_state.cv_data['skills']})
                PROFESSIONAL EXPERIENCE (Role | Company | Dates)
                User Data: {st.session_state.cv_data['raw_experience']}
                
                EDUCATION 
                - Degree: {st.session_state.cv_data.get('degree')}
                - University: {st.session_state.cv_data.get('university')}
                - College/Faculty: {st.session_state.cv_data.get('college')}
                - Year: {st.session_state.cv_data.get('grad_year')}
                
                {optional_prompt}
                LANGUAGES ({st.session_state.cv_data['languages']})
                """
                
                generated_text = safe_generate(prompt_cv)
                if "Error:" in generated_text: st.error(generated_text)
                else: st.session_state.final_cv = generated_text; st.rerun()

        if st.session_state.final_cv:
            st.text_area("Resume Editor", st.session_state.final_cv, height=500)
            c1, c2, c3 = st.columns(3)
            # زرار الـ PDF الآن يستخدم الدالة الجديدة التي تدعم العربي
            c1.download_button("⬇️ Download PDF", create_pdf(st.session_state.final_cv), file_name, "application/pdf")
            c2.download_button("⬇️ Download Word", create_docx(st.session_state.final_cv), word_file_name, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
            if c3.button("🔄 Regenerate"): st.session_state.final_cv = ""; st.rerun()
    
    with t2:
        if st.button("✨ Write Cover Letter"):
            with st.spinner("Writing..."):
                prompt_cl = f"Write a professional cover letter for {st.session_state.cv_data['name']} applying for {st.session_state.cv_data['target_title']}."
                st.session_state.cover_letter = safe_generate(prompt_cl)
                st.rerun()
        if st.session_state.cover_letter:
            st.text_area("Cover Letter", st.session_state.cover_letter, height=400)
            st.download_button("⬇️ Download Letter", create_docx(st.session_state.cover_letter), "Cover_Letter.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    
    with t3:
        if st.button("🔍 Check ATS Score"):
            with st.spinner("Analyzing..."):
                context = jd if jd else st.session_state.cv_data['target_title']
                prompt_ats = f"Analyze this CV against this Job/Role: {context}. Give a Score out of 100, list Missing Keywords, and suggest Improvements."
                st.session_state.ats_analysis = safe_generate(prompt_ats)
                st.rerun()
        if st.session_state.ats_analysis: st.write(st.session_state.ats_analysis)

    st.markdown("---")
    if st.button("Start Over"):
        st.session_state.step = 1; st.session_state.cv_data = {}; st.session_state.final_cv = ""; st.rerun()
